from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# Create document
doc = Document()

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Claude Skills Catalog')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('66 Skills for AI-Assisted Development')
run.font.size = Pt(16)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Version 1.0.0 | 2026\n')
run.font.size = Pt(12)
run = info.add_run('For Startups & Development Teams')
run.font.size = Pt(12)
run.italic = True

# Table of Contents placeholder
doc.add_page_break()
doc.add_heading('Table of Contents', 1)
doc.add_paragraph('1. Overview')
doc.add_paragraph('2. Claude Models Compatibility Matrix')
doc.add_paragraph('3. Skills by Category')
doc.add_paragraph('   3.1 Backend (6 skills)')
doc.add_paragraph('   3.2 Frontend (6 skills)')
doc.add_paragraph('   3.3 Mobile (4 skills)')
doc.add_paragraph('   3.4 DevOps (10 skills)')
doc.add_paragraph('   3.5 Data (6 skills)')
doc.add_paragraph('   3.6 AI (4 skills)')
doc.add_paragraph('   3.7 QA (10 skills)')
doc.add_paragraph('   3.8 Product (4 skills)')
doc.add_paragraph('   3.9 Security (4 skills)')
doc.add_paragraph('   3.10 Blockchain (3 skills)')
doc.add_paragraph('   3.11 Gamedev (3 skills)')
doc.add_paragraph('   3.12 IoT (3 skills)')
doc.add_paragraph('   3.13 Design (3 skills)')

# Overview
doc.add_page_break()
doc.add_heading('1. Overview', 1)
doc.add_paragraph('Claude Skills is a collection of 66 specialized skills for AI-assisted software development. Each skill provides structured instructions, code examples, and validation steps for Claude to follow.')

p = doc.add_paragraph('Key Benefits:')
p.add_run('\n• Ready-to-use code templates').bold = False
doc.add_paragraph('• Validated instructions for each domain')
doc.add_paragraph('• Model-specific recommendations')
doc.add_paragraph('• CI/CD integration ready')

# Models Matrix
doc.add_page_break()
doc.add_heading('2. Claude Models Compatibility Matrix', 1)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Speed'
hdr_cells[2].text = 'Context'
hdr_cells[3].text = 'Best For'

data = [
    ['Haiku', '⚡⚡⚡', '200K', 'Simple, repetitive tasks'],
    ['Sonnet', '⚡⚡', '200K', 'Balanced - most skills'],
    ['Opus', '⚡', '200K', 'Complex reasoning, architecture']
]
for i, row_data in enumerate(data, 1):
    cells = table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph('\nRecommendation:')
doc.add_paragraph('• For most users: Sonnet (best balance of performance and cost)')
doc.add_paragraph('• For complex tasks: Opus (architecture, AI agents)')
doc.add_paragraph('• For simple tasks: Haiku (sufficient and fast)')

# Skills by Category
doc.add_page_break()
doc.add_heading('3. Skills by Category', 1)

# Load skills data
with open('skills_catalog.json', 'r', encoding='utf-8-sig') as f:
    data = json.load(f)

skills_by_cat = {}
for skill in data['skills']:
    cat = skill['category']
    if cat not in skills_by_cat:
        skills_by_cat[cat] = []
    skills_by_cat[cat].append(skill)

# Category display names
cat_names = {
    'backend': 'Backend',
    'frontend': 'Frontend',
    'mobile': 'Mobile',
    'devops': 'DevOps',
    'data': 'Data',
    'ai': 'AI',
    'qa': 'QA',
    'product': 'Product',
    'security': 'Security',
    'block': 'Blockchain',
    'gamedev': 'GameDev',
    'iot': 'IoT',
    'design': 'Design'
}

for i, (cat, skills) in enumerate(sorted(skills_by_cat.items()), 3):
    doc.add_heading(f'{i}. {cat_names.get(cat, cat)} ({len(skills)} skills)', 2)
    
    # Create table for skills
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Skill'
    hdr[1].text = 'Description'
    hdr[2].text = 'Models'
    hdr[3].text = 'Tags'
    
    for skill in skills:
        row_cells = table.add_row().cells
        row_cells[0].text = skill['name']
        row_cells[1].text = skill['description'][:80] + '...' if len(skill['description']) > 80 else skill['description']
        row_cells[2].text = ', '.join(skill['models'])
        row_cells[3].text = ', '.join(skill['tags'][:3])
    
    doc.add_paragraph()

# Save
doc.save('D:/claude-ai-skills/Claude_Skills_Catalog_Startup.docx')
print('Document created: Claude_Skills_Catalog_Startup.docx')
